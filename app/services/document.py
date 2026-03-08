"""
NAYAM (नयम्) — Document Service.

Business logic for document management including
secure file upload, real text extraction, AI summarization,
and RAG embedding storage.
"""

import logging
import os
import uuid
from typing import List, Tuple
from uuid import UUID

from fastapi import HTTPException, UploadFile, status
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.document import Document
from app.repositories.document import DocumentRepository

logger = logging.getLogger(__name__)

settings = get_settings()

# ── Allowed File Extensions (security) ───────────────────────────────
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".doc", ".docx", ".png", ".jpg", ".jpeg"}


def _validate_file(file: UploadFile) -> None:
    """
    Validate an uploaded file for security.
    """
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Filename is required.",
        )

    _, ext = os.path.splitext(file.filename)
    if ext.lower() not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type '{ext}' not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )


async def _save_file(file: UploadFile) -> str:
    """
    Save an uploaded file to the configured upload directory.
    """
    _, ext = os.path.splitext(file.filename or "file")
    safe_filename = f"{uuid.uuid4().hex}{ext.lower()}"
    file_path = os.path.join(settings.UPLOAD_DIR, safe_filename)

    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    content = await file.read()

    if len(content) > settings.max_upload_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds maximum size of {settings.MAX_UPLOAD_SIZE_MB} MB.",
        )

    with open(file_path, "wb") as f:
        f.write(content)

    logger.info("File saved: %s (%d bytes)", file_path, len(content))
    return file_path


# ── Text Extraction ──────────────────────────────────────────────────

def extract_text(file_path: str) -> str:
    """
    Extract text from a document file.

    Supports: .txt, .pdf, .docx
    Other formats return a placeholder.
    """
    ext = os.path.splitext(file_path)[1].lower()
    logger.info("Extracting text from %s (type: %s)", file_path, ext)

    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        elif ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n\n".join(text_parts) if text_parts else "[No text could be extracted from PDF]"
            except Exception as e:
                logger.warning("PDF extraction failed: %s", e)
                return f"[PDF text extraction failed: {e}]"

        elif ext in (".doc", ".docx"):
            try:
                import docx
                doc = docx.Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs) if paragraphs else "[No text found in document]"
            except Exception as e:
                logger.warning("DOCX extraction failed: %s", e)
                return f"[DOCX text extraction failed: {e}]"

        else:
            return f"[Text extraction not supported for {ext} files]"

    except Exception as e:
        logger.error("Text extraction failed for %s: %s", file_path, e)
        return f"[Text extraction failed: {e}]"


def chunk_text(text: str, chunk_size: int = 400, overlap: int = 50) -> List[str]:
    """
    Split text into overlapping chunks suitable for RAG retrieval.

    Args:
        text: The full document text.
        chunk_size: Number of words per chunk.
        overlap: Word overlap between chunks.

    Returns:
        List of text chunk strings.
    """
    words = text.split()
    if not words:
        return []

    if len(words) <= chunk_size:
        return [text.strip()]

    chunks = []
    step = max(1, chunk_size - overlap)
    for i in range(0, len(words), step):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk.strip())
        if i + chunk_size >= len(words):
            break

    return chunks


def generate_summary(text: str) -> str:
    """
    Generate an AI summary of the document text using Groq.

    Falls back to a simple extractive summary if LLM is unavailable.
    """
    # Try LLM summary
    try:
        from app.core.config import get_settings
        s = get_settings()
        if s.GROQ_API_KEY:
            from groq import Groq
            client = Groq(api_key=s.GROQ_API_KEY)

            # Use first ~2000 words for summary
            truncated = " ".join(text.split()[:2000])

            response = client.chat.completions.create(
                model=s.GROQ_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a document summarizer for the NAYAM governance platform. "
                            "Provide a concise 2-3 sentence summary of the document. "
                            "Focus on key policy points, action items, and relevance to "
                            "municipal governance."
                        ),
                    },
                    {
                        "role": "user",
                        "content": f"Summarize this document:\n\n{truncated}",
                    },
                ],
                temperature=0.3,
                max_tokens=200,
            )
            return response.choices[0].message.content
    except Exception as e:
        logger.warning("LLM summary generation failed: %s", e)

    # Fallback: extractive summary (first 3 sentences)
    sentences = text.replace("\n", " ").split(".")
    summary_sentences = [s.strip() for s in sentences[:3] if s.strip()]
    if summary_sentences:
        return ". ".join(summary_sentences) + "."
    return "Document uploaded successfully. Summary generation pending."


class DocumentService:
    """
    Service layer for document operations.
    """

    def __init__(self, db: Session) -> None:
        self.repo = DocumentRepository(db)
        self._db = db

    async def upload_document(
        self,
        title: str,
        file: UploadFile,
        uploaded_by: UUID,
    ) -> Document:
        """
        Upload a document with real text extraction, AI summary, and RAG embedding.
        """
        _validate_file(file)
        file_path = await _save_file(file)

        # Real text extraction
        extracted_text = extract_text(file_path)

        # AI-powered summary
        summary = generate_summary(extracted_text)

        document = Document(
            title=title,
            uploaded_by=uploaded_by,
            file_path=file_path,
            extracted_text=extracted_text,
            summary=summary,
        )
        created_doc = self.repo.create(document)

        # Store text chunks as embeddings for RAG retrieval
        try:
            self._store_document_embeddings(created_doc)
        except Exception as e:
            logger.warning("Failed to store embeddings for document %s: %s", created_doc.id, e)

        return created_doc

    def _store_document_embeddings(self, document: Document) -> None:
        """
        Chunk the document text and store each chunk with dense FAISS embeddings.
        """
        from app.services.memory import MemoryService, generate_embeddings_batch

        if not document.extracted_text or document.extracted_text.startswith("["):
            return

        memory = MemoryService(self._db)
        chunks = chunk_text(document.extracted_text)

        if not chunks:
            return

        # Batch-encode all chunks at once for efficiency
        vectors = generate_embeddings_batch(chunks)

        for idx, (chunk, vec) in enumerate(zip(chunks, vectors)):
            memory.store_embedding(
                source_type="document",
                source_id=document.id,
                text=chunk,
                embedding_vector=vec,
                chunk_index=idx,
                model_name="all-MiniLM-L6-v2",
            )

        logger.info(
            "Stored %d text chunks for document %s (%s)",
            len(chunks), document.id, document.title,
        )

    def get_document(self, doc_id: UUID) -> Document:
        """Retrieve a single document by ID."""
        document = self.repo.get_by_id(doc_id)
        if document is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document with id {doc_id} not found.",
            )
        return document

    def list_documents(
        self,
        skip: int = 0,
        limit: int = 50,
    ) -> Tuple[List[Document], int]:
        """List documents with pagination."""
        return self.repo.get_all(skip=skip, limit=limit)

    def delete_document(self, doc_id: UUID) -> None:
        """Delete a document and optionally remove the file."""
        document = self.get_document(doc_id)

        if document.file_path and os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
                logger.info("Removed file: %s", document.file_path)
            except OSError as e:
                logger.warning("Failed to remove file %s: %s", document.file_path, e)

        self.repo.delete(document)
        logger.info("Document deleted: %s", doc_id)
